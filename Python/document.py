from docx import Document

# Tạo một đối tượng Document mới
document = Document()

# Tiêu đề chính
document.add_heading('Tài liệu Thiết kế Kỹ thuật: RESTful API Login', 0)

# 1. Giới thiệu
document.add_heading('1. Giới thiệu', level=1)
document.add_paragraph(
    "Tài liệu này trình bày chi tiết thiết kế kỹ thuật cho tính năng đăng nhập (Login) sử dụng RESTful API của hệ thống. "
    "Mục tiêu của tính năng là:\n"
    "- Xác thực người dùng với quy trình bảo mật cao.\n"
    "- Hỗ trợ đăng nhập cho Admin: sử dụng email và password.\n"
    "- Hỗ trợ đăng nhập cho Customer: sử dụng email, password và tích hợp đăng nhập qua Google, Facebook."
)

# 2. Phạm vi
document.add_heading('2. Phạm vi', level=1)
document.add_paragraph(
    "Tính năng đăng nhập này áp dụng cho:\n"
    "- Admin: Nhân viên quản trị hệ thống.\n"
    "- Customer: Người dùng cuối của hệ thống."
)

# 3. Yêu cầu chức năng
document.add_heading('3. Yêu cầu chức năng', level=1)

# 3.1 Yêu cầu chung
document.add_heading('3.1 Yêu cầu chung', level=2)
document.add_paragraph(
    "Hệ thống phải xác thực và phân quyền người dùng dựa trên thông tin đăng nhập. Sinh ra token (ví dụ: JWT) sau khi đăng nhập thành công để thực hiện các giao tiếp sau này. "
    "Hỗ trợ xác thực OAuth cho các nhà cung cấp Google và Facebook. Các dữ liệu nhạy cảm (như mật khẩu) phải được mã hóa theo tiêu chuẩn bảo mật."
)

# 3.2 Yêu cầu đối với Admin
document.add_heading('3.2 Yêu cầu đối với Admin', level=2)
document.add_paragraph(
    "- Admin đăng nhập bằng email và password.\n"
    "- Kiểm tra tính hợp lệ của email và password.\n"
    "- Trả về thông tin admin cùng với token khi đăng nhập thành công."
)

# 3.3 Yêu cầu đối với Customer
document.add_heading('3.3 Yêu cầu đối với Customer', level=2)
document.add_paragraph(
    "- Đăng nhập bằng email và password cho tài khoản nội bộ.\n"
    "- Đăng nhập qua Google và Facebook:\n"
    "   - Xác thực token được cấp từ nhà cung cấp OAuth.\n"
    "   - Nếu người dùng chưa tồn tại, thực hiện đăng ký tự động.\n"
    "   - Trả về token và thông tin khách hàng khi đăng nhập thành công."
)

# 4. Kiến trúc hệ thống
document.add_heading('4. Kiến trúc hệ thống', level=1)

# 4.1 Thành phần chính
document.add_heading('4.1 Thành phần chính', level=2)
document.add_paragraph(
    "- API Gateway: Tiếp nhận và định tuyến các yêu cầu từ client.\n"
    "- Authentication Service: Xử lý đăng nhập, xác thực và sinh token.\n"
    "- User Service: Quản lý dữ liệu người dùng (Admin và Customer).\n"
    "- OAuth Integration Module: Tích hợp với Google và Facebook để xác thực người dùng đăng nhập qua mạng xã hội."
)

# 4.2 Luồng xử lý
document.add_heading('4.2 Luồng xử lý', level=2)

# 4.2.1 Đăng nhập Admin
document.add_heading('4.2.1 Đăng nhập Admin', level=3)
document.add_paragraph(
    "1. Admin gửi yêu cầu POST đến `/api/admin/login` với email và password.\n"
    "2. API Gateway chuyển tiếp yêu cầu đến Authentication Service.\n"
    "3. Authentication Service xác thực thông tin với cơ sở dữ liệu.\n"
    "4. Nếu hợp lệ, sinh JWT token và trả về cho client.\n"
    "5. Nếu không hợp lệ, trả về lỗi 401 Unauthorized."
)

# 4.2.2 Đăng nhập Customer qua email/password
document.add_heading('4.2.2 Đăng nhập Customer qua email/password', level=3)
document.add_paragraph(
    "1. Customer gửi yêu cầu POST đến `/api/customer/login` với email và password.\n"
    "2. Authentication Service xác thực thông tin người dùng.\n"
    "3. Nếu hợp lệ, sinh JWT token và trả về cho client.\n"
    "4. Nếu không hợp lệ, trả về lỗi thích hợp."
)

# 4.2.3 Đăng nhập Customer qua Google/Facebook
document.add_heading('4.2.3 Đăng nhập Customer qua Google/Facebook', level=3)
document.add_paragraph(
    "1. Customer gửi yêu cầu POST đến `/api/customer/login` với các trường:\n"
    "   - provider: \"google\" hoặc \"facebook\"\n"
    "   - token: token OAuth được cấp từ nhà cung cấp.\n"
    "2. OAuth Integration Module xác thực token với Google hoặc Facebook.\n"
    "3. Nếu token hợp lệ:\n"
    "   - Nếu người dùng đã tồn tại: tiến hành đăng nhập và sinh token nội bộ.\n"
    "   - Nếu người dùng chưa tồn tại: thực hiện đăng ký người dùng, sau đó sinh token.\n"
    "4. Trả về token cùng thông tin người dùng."
)

# 5. Chi tiết API
document.add_heading('5. Chi tiết API', level=1)

# 5.1 API cho Admin
document.add_heading('5.1 API cho Admin', level=2)
document.add_paragraph("URL: `/api/admin/login`")
document.add_paragraph("Method: POST")
document.add_paragraph("Headers: `Content-Type: application/json`")
document.add_paragraph("Request Body:")
document.add_paragraph(
    '{\n  "email": "admin@example.com",\n  "password": "password123"\n}',
    style='Intense Quote'
)
document.add_paragraph(
    "Success Response:\nStatus Code: 200 OK\nResponse Body:\n{\n  \"accessToken\": \"jwt_token_here\",\n  \"refreshToken\": \"refresh_token_here\",\n  \"admin\": {\n    \"id\": \"admin_id\",\n    \"email\": \"admin@example.com\",\n    \"name\": \"Admin Name\"\n  }\n}"
)
document.add_paragraph(
    "Error Responses:\nStatus Code: 401 Unauthorized\nResponse Body:\n{\n  \"error\": \"Invalid credentials\"\n}\nStatus Code: 400 Bad Request (nếu thiếu thông tin bắt buộc)"
)

# 5.2 API cho Customer
document.add_heading('5.2 API cho Customer', level=2)

# 5.2.1 Đăng nhập qua Email/Password
document.add_heading('5.2.1 Đăng nhập qua Email/Password', level=3)
document.add_paragraph("URL: `/api/customer/login`")
document.add_paragraph("Method: POST")
document.add_paragraph("Headers: `Content-Type: application/json`")
document.add_paragraph("Request Body:")
document.add_paragraph(
    '{\n  "email": "customer@example.com",\n  "password": "password123"\n}',
    style='Intense Quote'
)
document.add_paragraph(
    "Success Response:\nStatus Code: 200 OK\nResponse Body:\n{\n  \"accessToken\": \"jwt_token_here\",\n  \"refreshToken\": \"refresh_token_here\",\n  \"customer\": {\n    \"id\": \"customer_id\",\n    \"email\": \"customer@example.com\",\n    \"name\": \"Customer Name\"\n  }\n}"
)
document.add_paragraph(
    "Error Responses:\nStatus Code: 401 Unauthorized\nResponse Body:\n{\n  \"error\": \"Invalid credentials\"\n}\nStatus Code: 400 Bad Request"
)

# 5.2.2 Đăng nhập qua Google/Facebook
document.add_heading('5.2.2 Đăng nhập qua Google/Facebook', level=3)
document.add_paragraph("URL: `/api/customer/login`")
document.add_paragraph("Method: POST")
document.add_paragraph("Headers: `Content-Type: application/json`")
document.add_paragraph("Request Body (ví dụ với Google):")
document.add_paragraph(
    '{\n  "provider": "google",\n  "token": "google_oauth_token_here"\n}',
    style='Intense Quote'
)
document.add_paragraph("Request Body (ví dụ với Facebook):")
document.add_paragraph(
    '{\n  "provider": "facebook",\n  "token": "facebook_oauth_token_here"\n}',
    style='Intense Quote'
)
document.add_paragraph(
    "Success Response:\nStatus Code: 200 OK\nResponse Body:\n{\n  \"accessToken\": \"jwt_token_here\",\n  \"refreshToken\": \"refresh_token_here\",\n  \"customer\": {\n    \"id\": \"customer_id\",\n    \"email\": \"customer@example.com\",\n    \"name\": \"Customer Name\",\n    \"provider\": \"google\" // hoặc \"facebook\"\n  }\n}"
)
document.add_paragraph(
    "Error Responses:\nStatus Code: 401 Unauthorized\nResponse Body:\n{\n  \"error\": \"Invalid or expired OAuth token\"\n}\nStatus Code: 400 Bad Request"
)

# 6. Mô hình dữ liệu
document.add_heading('6. Mô hình dữ liệu', level=1)

# 6.1 Bảng Admin
document.add_heading('6.1 Bảng Admin', level=2)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Trường'
hdr_cells[1].text = 'Kiểu dữ liệu'
hdr_cells[2].text = 'Mô tả'
row_cells = table.add_row().cells
row_cells[0].text = 'id'
row_cells[1].text = 'UUID'
row_cells[2].text = 'Mã định danh admin'
row_cells = table.add_row().cells
row_cells[0].text = 'email'
row_cells[1].text = 'String'
row_cells[2].text = 'Email đăng nhập'
row_cells = table.add_row().cells
row_cells[0].text = 'password'
row_cells[1].text = 'String'
row_cells[2].text = 'Mật khẩu đã mã hóa'
row_cells = table.add_row().cells
row_cells[0].text = 'name'
row_cells[1].text = 'String'
row_cells[2].text = 'Tên admin'
row_cells = table.add_row().cells
row_cells[0].text = 'created_at'
row_cells[1].text = 'Timestamp'
row_cells[2].text = 'Thời gian tạo'
row_cells = table.add_row().cells
row_cells[0].text = 'updated_at'
row_cells[1].text = 'Timestamp'
row_cells[2].text = 'Thời gian cập nhật'

# 6.2 Bảng Customer
document.add_heading('6.2 Bảng Customer', level=2)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Trường'
hdr_cells[1].text = 'Kiểu dữ liệu'
hdr_cells[2].text = 'Mô tả'
row_cells = table.add_row().cells
row_cells[0].text = 'id'
row_cells[1].text = 'UUID'
row_cells[2].text = 'Mã định danh khách hàng'
row_cells = table.add_row().cells
row_cells[0].text = 'email'
row_cells[1].text = 'String'
row_cells[2].text = 'Email khách hàng'
row_cells = table.add_row().cells
row_cells[0].text = 'password'
row_cells[1].text = 'String'
row_cells[2].text = 'Mật khẩu (có thể null nếu đăng nhập social)'
row_cells = table.add_row().cells
row_cells[0].text = 'name'
row_cells[1].text = 'String'
row_cells[2].text = 'Tên khách hàng'
row_cells = table.add_row().cells
row_cells[0].text = 'provider'
row_cells[1].text = 'String'
row_cells[2].text = 'Nguồn đăng nhập: local, google, facebook'
row_cells = table.add_row().cells
row_cells[0].text = 'created_at'
row_cells[1].text = 'Timestamp'
row_cells[2].text = 'Thời gian tạo'
row_cells = table.add_row().cells
row_cells[0].text = 'updated_at'
row_cells[1].text = 'Timestamp'
row_cells[2].text = 'Thời gian cập nhật'

# 7. Bảo mật và xác thực
document.add_heading('7. Bảo mật và xác thực', level=1)
document.add_paragraph(
    "- Mã hóa mật khẩu: Sử dụng bcrypt hoặc các thuật toán tương đương.\n"
    "- JWT Token: Dùng JSON Web Token (JWT) cho các phiên đăng nhập.\n"
    "- Xác thực OAuth: Tích hợp xác thực qua API của Google và Facebook để kiểm tra token.\n"
    "- HTTPS: Bắt buộc tất cả giao tiếp giữa client và server thông qua HTTPS."
)

# 8. Xử lý lỗi
document.add_heading('8. Xử lý lỗi', level=1)
document.add_paragraph(
    "Các lỗi được xử lý theo mã trạng thái HTTP:\n"
    "- 400 Bad Request: Yêu cầu không hợp lệ hoặc thiếu dữ liệu cần thiết.\n"
    "- 401 Unauthorized: Thông tin đăng nhập sai hoặc token không hợp lệ.\n"
    "- 500 Internal Server Error: Lỗi phía máy chủ.\n"
    "Mỗi lỗi phải trả về thông tin chi tiết dưới dạng JSON để client có thể xử lý."
)

# 9. Kiểm thử và đảm bảo chất lượng
document.add_heading('9. Kiểm thử và đảm bảo chất lượng', level=1)
document.add_paragraph(
    "- Unit Tests: Kiểm thử từng hàm xử lý đăng nhập, xác thực và sinh token.\n"
    "- Integration Tests: Kiểm thử toàn bộ luồng đăng nhập từ client đến server.\n"
    "- Security Tests: Thực hiện kiểm thử bảo mật như SQL Injection, XSS, CSRF, và brute force."
)

# 10. Ghi chú và mở rộng
document.add_heading('10. Ghi chú và mở rộng', level=1)
document.add_paragraph(
    "- Hệ thống cần linh hoạt mở rộng cho các phương thức đăng nhập khác (ví dụ: Twitter, LinkedIn).\n"
    "- Định kỳ cập nhật các thuật toán mã hóa và quy trình bảo mật theo tiêu chuẩn mới.\n"
    "- Ghi log chi tiết các lần đăng nhập thất bại để hỗ trợ việc giám sát và xử lý vi phạm bảo mật."
)

# 11. Phụ lục
document.add_heading('11. Phụ lục', level=1)

# 11.1 Quy trình đăng nhập qua OAuth
document.add_heading('11.1 Quy trình đăng nhập qua OAuth', level=2)
document.add_paragraph(
    "1. Client lấy token từ nhà cung cấp (Google/Facebook).\n"
    "2. Client gửi token này kèm theo thông tin `provider` đến API của hệ thống.\n"
    "3. Hệ thống xác thực token với API của Google hoặc Facebook.\n"
    "4. Nếu xác thực thành công:\n"
    "   - Kiểm tra xem người dùng đã tồn tại chưa.\n"
    "   - Nếu chưa, tạo mới tài khoản khách hàng.\n"
    "   - Sinh JWT token nội bộ và trả về cho client.\n"
    "5. Nếu xác thực không thành công, trả về lỗi 401 Unauthorized."
)

# 11.2 Ví dụ mã code (Pseudo-code)
document.add_heading('11.2 Ví dụ mã code (Pseudo-code)', level=2)
document.add_paragraph(
    "def login_admin(email, password):\n"
    "    admin = find_admin_by_email(email)\n"
    "    if admin and verify_password(password, admin.password):\n"
    "        token = generate_jwt(admin.id)\n"
    "        return {\"accessToken\": token, \"admin\": admin}\n"
    "    else:\n"
    "        raise UnauthorizedError(\"Invalid credentials\")\n\n"
    "def login_customer(email=None, password=None, provider=None, oauth_token=None):\n"
    "    if provider in [\"google\", \"facebook\"]:\n"
    "        user_info = verify_oauth_token(provider, oauth_token)\n"
    "        customer = find_or_create_customer(user_info, provider)\n"
    "    else:\n"
    "        customer = find_customer_by_email(email)\n"
    "        if not customer or not verify_password(password, customer.password):\n"
    "            raise UnauthorizedError(\"Invalid credentials\")\n"
    "    token = generate_jwt(customer.id)\n"
    "    return {\"accessToken\": token, \"customer\": customer}"
)

# Lưu file
file_name = "technical_design_document.docx"
document.save(file_name)
print(f"File '{file_name}' created successfully!")

